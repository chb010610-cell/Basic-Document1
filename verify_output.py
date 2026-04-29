from docx import Document
import os

output_dir = r"C:\Users\Administrator\WorkBuddy\20260429103845\output"

# Check each output file for correct replacement
test_company = "宁波市金狼电器有限公司"
old_company = "宁波市海曙新艺洁具有限公司"

for fname in os.listdir(output_dir):
    if fname.endswith('.docx'):
        fpath = os.path.join(output_dir, fname)
        try:
            doc = Document(fpath)
            new_count = 0
            old_count = 0
            for para in doc.paragraphs:
                if test_company in para.text:
                    new_count += 1
                if old_company in para.text:
                    old_count += 1
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if test_company in cell.text:
                            new_count += 1
                        if old_company in cell.text:
                            old_count += 1
            print(f"{fname}: NEW={new_count}, OLD(remaining)={old_count}")
        except Exception as e:
            print(f"{fname}: ERROR {e}")