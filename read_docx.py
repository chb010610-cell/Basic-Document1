from docx import Document
import json
import re
import sys

files = [
    r"C:\Users\Administrator\Desktop\后台\1程序文件封面.docx",
    r"C:\Users\Administrator\Desktop\后台\2质量手册.docx",
    r"C:\Users\Administrator\Desktop\后台\3.程序.docx",
    r"C:\Users\Administrator\Desktop\后台\4管理手册.docx",
    r"C:\Users\Administrator\Desktop\后台\5受控文件清单.docx",
    r"C:\Users\Administrator\Desktop\后台\7文件发放、收回和销毁记录.docx",
    r"C:\Users\Administrator\Desktop\后台\9-岗位说明书.docx",
    r"C:\Users\Administrator\Desktop\后台\10质量各种任命通知(3).docx",
]

all_placeholders = set()

for fpath in files:
    try:
        doc = Document(fpath)
        fname = fpath.split("\\")[-1]
        print(f"\n=== {fname} ===")
        file_placeholders = set()
        
        for para in doc.paragraphs:
            text = para.text
            # Find placeholders like {{xxx}} or {xxx} or 【xxx】
            matches = re.findall(r'[{【].*?[}】]', text)
            for m in matches:
                file_placeholders.add(m)
                all_placeholders.add(m)
        
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    matches = re.findall(r'[{【].*?[}】]', text)
                    for m in matches:
                        file_placeholders.add(m)
                        all_placeholders.add(m)
        
        for p in sorted(file_placeholders):
            print(f"  {p}")
    except Exception as e:
        print(f"Error reading {fpath}: {e}")

print("\n\n=== ALL UNIQUE PLACEHOLDERS ===")
for p in sorted(all_placeholders):
    print(p)