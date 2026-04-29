from docx import Document
import json

# Look at all docx files for actual company-specific content
files = [
    ("1程序文件封面.docx", r"C:\Users\Administrator\Desktop\后台\1程序文件封面.docx"),
    ("2质量手册.docx", r"C:\Users\Administrator\Desktop\后台\2质量手册.docx"),
    ("3.程序.docx", r"C:\Users\Administrator\Desktop\后台\3.程序.docx"),
    ("4管理手册.docx", r"C:\Users\Administrator\Desktop\后台\4管理手册.docx"),
    ("5受控文件清单.docx", r"C:\Users\Administrator\Desktop\后台\5受控文件清单.docx"),
    ("7文件发放收回销毁记录.docx", r"C:\Users\Administrator\Desktop\后台\7文件发放、收回和销毁记录.docx"),
    ("9岗位说明书.docx", r"C:\Users\Administrator\Desktop\后台\9-岗位说明书.docx"),
    ("10任命通知.docx", r"C:\Users\Administrator\Desktop\后台\10质量各种任命通知(3).docx"),
]

# Search for known company names in the docs
company_names = ["宁波市海曙新艺洁具有限公司", "宁波市金狼电器有限公司", "夏亚明", "龚伟", "贾宇", "XY"]

for fname, fpath in files:
    try:
        doc = Document(fpath)
        found = {}
        for search in company_names:
            count = 0
            for para in doc.paragraphs:
                if search in para.text:
                    count += 1
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if search in cell.text:
                            count += 1
            if count > 0:
                found[search] = count
        
        if found:
            print(f"{fname}: {json.dumps(found, ensure_ascii=False)}")
        else:
            print(f"{fname}: no matches")
    except Exception as e:
        print(f"{fname}: ERROR {e}")