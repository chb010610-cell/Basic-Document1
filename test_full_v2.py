# -*- coding: utf-8 -*-
import sys, os, io, json, urllib.request
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.path.insert(0, r'C:\Users\Administrator\.workbuddy\binaries\python\envs\default\Lib\site-packages')

data = {
    "company_name": "温州测试电器有限公司",
    "address": "浙江省温州市鹿城区测试路88号",
    "product": "插座、开关等电器配件",
    "english_abbr": "CS",
    "doc_date": "2026-07-01",
    "version_no": "C/0",
    "gm_name": "王测试",
    "production_mgr": "李生产",
    "mgr_rep": "张管理",
    "hr_head": "赵人事",
    "sales_head": "钱业务",
    "purchase_head": "孙采购",
    "tech_head": "周技术",
    "raw_warehouse": "吴仓库",
    "finished_warehouse": "郑成品",
    "mechanic": "冯机修",
    "iqc_person": "陈进料",
    "ipqc1_person": "褚制程1",
    "ipqc2_person": "卫制程2",
    "ipqc3_person": "蒋制程3",
    "fqc_person": "沈成品",
    "supplier": "测试供应商有限公司",
    "license_code": "91330300MA5TEST02"
}

url = "http://127.0.0.1:8765/api/generate"
req = urllib.request.Request(url, data=json.dumps(data, ensure_ascii=False).encode('utf-8'),
                              headers={'Content-Type': 'application/json; charset=utf-8'})
print("Sending request...")
resp = urllib.request.urlopen(req, timeout=180)
result = json.loads(resp.read().decode('utf-8'))
print(f"Success: {result['success']}  Count: {result['count']}")
total_repl = 0
for g in result.get('generated', []):
    repl = g.get('replacements', 0)
    err = g.get('error', '')
    total_repl += repl
    if err:
        print(f"  [FAIL] {g['name']}: {err}")
    else:
        print(f"  [OK] {g['name']}: {repl} replacements")
print(f"\nTotal replacements across all docs: {total_repl}")

# Verify header replacement in 3.程序.docx
from docx import Document
doc3_path = r'C:\Users\Administrator\WorkBuddy\20260429103845\output\3.程序.docx'
if os.path.exists(doc3_path):
    doc = Document(doc3_path)
    print("\n=== Verify 3.程序.docx headers ===")
    header_count = 0
    for section in doc.sections:
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    ct = cell.text.strip()
                    if ct and ('温州' in ct or '2026-07' in ct or 'C/0' in ct):
                        print(f"  [OK-REPLACED] {ct[:80]}")
                        header_count += 1
                    elif ct and ('宁波' in ct or '2025-03' in ct or 'A/0' == ct or 'B/0' == ct):
                        print(f"  [!! NOT REPLACED] {ct[:80]}")
                        header_count += 1
    if header_count == 0:
        print("  No matching header content found")

# Verify 李忠伟 replacement
exam_path = r'C:\Users\Administrator\WorkBuddy\20260429103845\output\检验人员培训考试.docx'
if os.path.exists(exam_path):
    doc = Document(exam_path)
    print("\n=== Verify 检验人员培训考试.docx ===")
    found_lzw = False
    found_iqc = False
    for para in doc.paragraphs:
        if '李忠伟' in para.text:
            found_lzw = True
            print(f"  [!!] Still has 李忠伟: {para.text.strip()[:80]}")
        if '陈进料' in para.text:
            found_iqc = True
            print(f"  [OK] Replaced with 陈进料: {para.text.strip()[:80]}")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '李忠伟' in cell.text:
                    found_lzw = True
                    print(f"  [!!] Still has 李忠伟 (table): {cell.text.strip()[:80]}")
                if '陈进料' in cell.text:
                    found_iqc = True
                    print(f"  [OK] Replaced with 陈进料 (table): {cell.text.strip()[:80]}")
    if not found_lzw:
        print("  [OK] 李忠伟 fully replaced!")
    if not found_iqc:
        print("  [!!] 陈进料 not found in output")
