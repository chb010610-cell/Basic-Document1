from docx import Document
import json
import re

# Read the quality manual to find actual company names and replaceable text
fpath = r"C:\Users\Administrator\Desktop\后台\2质量手册.docx"
doc = Document(fpath)

# Sample first 100 paragraphs
print("=== Paragraphs (first 100) ===")
for i, para in enumerate(doc.paragraphs[:100]):
    if para.text.strip():
        print(f"  [{i}] {para.text[:200]}")

# Sample first 5 tables
print("\n=== Tables (first 5) ===")
for ti, table in enumerate(doc.tables[:5]):
    print(f"\n--- Table {ti} ---")
    for ri, row in enumerate(table.rows[:10]):
        cells = [cell.text[:50] for cell in row.cells]
        print(f"  Row {ri}: {json.dumps(cells, ensure_ascii=False)}")