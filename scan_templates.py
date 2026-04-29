# -*- coding: utf-8 -*-
"""扫描模板中的替换值 - v2"""
import sys, os, re, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.path.insert(0, r'C:\Users\Administrator\.workbuddy\binaries\python\envs\default\Lib\site-packages')
from docx import Document
from collections import Counter

template_dir = r'C:\Users\Administrator\Desktop\website后台'
all_texts = Counter()

for fn in sorted(os.listdir(template_dir)):
    if not fn.endswith('.docx'):
        continue
    fp = os.path.join(template_dir, fn)
    try:
        doc = Document(fp)
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                all_texts[t] += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        all_texts[t] += 1
    except Exception as e:
        print(f'Error: {fn}: {e}')

# 统计关键值
targets = ['2025-03-05', 'B/0', '夏亚明', '宁波市海曙新艺洁具有限公司', '浙江省宁波市海曙区横街镇桃源村',
           '拖把、刷子、扫把等日用洁具', '龚伟', '贾宇', '于静洁', '项万青', '阮璐', '王波',
           '张丽容', '俄底只桑', '夏方强', '戴海能', '解欧红', '李丹',
           '浙江旭华包装有限公司', '91330212X10144809X']

print("=== KEY VALUE FREQUENCY ===")
for name in targets:
    total = sum(c for t, c in all_texts.items() if name in t)
    print(f"  [{total:>5}x] {name}")

# 查找所有包含日期的文本
print("\n=== DATE TEXTS ===")
for t, c in sorted(all_texts.items()):
    if ('2025' in t or '2026' in t or '2024' in t) and len(t) < 60:
        print(f"  [{c}x] {t.replace(chr(10), ' | ')}")

# 查找包含B/0的
print("\n=== VERSION TEXTS ===")
for t, c in sorted(all_texts.items()):
    if ('B/0' in t) and len(t) < 60:
        print(f"  [{c}x] {t.replace(chr(10), ' | ')}")

# 检查公司名称的变体形式
print("\n=== COMPANY NAME VARIANTS ===")
for t, c in sorted(all_texts.items()):
    if '新艺洁具' in t or '海曙新艺' in t:
        print(f"  [{c}x] {t[:80].replace(chr(10), ' | ')}")

print("\nDone.")
