# -*- coding: utf-8 -*-
import sys, os, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.path.insert(0, r'C:\Users\Administrator\.workbuddy\binaries\python\envs\default\Lib\site-packages')
from docx import Document
from docx.oxml.ns import qn

template_dir = r'C:\Users\Administrator\Desktop\website后台'

print("=== 1. 扫描所有文件的页眉/页脚内容 ===")
for fn in sorted(os.listdir(template_dir)):
    if not fn.endswith('.docx'):
        continue
    fp = os.path.join(template_dir, fn)
    try:
        doc = Document(fp)
        headers_found = []
        footers_found = []

        for section in doc.sections:
            # 页眉
            header = section.header
            for para in header.paragraphs:
                t = para.text.strip()
                if t:
                    headers_found.append(t[:100])
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t:
                            headers_found.append(f"[TABLE] {t[:100]}")

            # 页脚
            footer = section.footer
            for para in footer.paragraphs:
                t = para.text.strip()
                if t:
                    footers_found.append(t[:100])
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t:
                            footers_found.append(f"[TABLE] {t[:100]}")

        if headers_found or footers_found:
            print(f"\n--- {fn} ---")
            for h in headers_found:
                print(f"  HEADER: {h}")
            for f in footers_found:
                print(f"  FOOTER: {f}")
        else:
            print(f"{fn}: (no header/footer text)")
    except Exception as e:
        print(f"Error {fn}: {e}")

print("\n\n=== 2. 扫描'李忠伟'在哪些文件中出现 ===")
for fn in sorted(os.listdir(template_dir)):
    if not fn.endswith('.docx'):
        continue
    fp = os.path.join(template_dir, fn)
    try:
        doc = Document(fp)
        count = 0
        examples = []
        for para in doc.paragraphs:
            if '李忠伟' in para.text:
                count += 1
                if len(examples) < 3:
                    examples.append(para.text.strip()[:80])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '李忠伟' in cell.text:
                        count += 1
                        if len(examples) < 3:
                            examples.append(cell.text.strip()[:80])

        # 也检查页眉页脚中的李忠伟
        for section in doc.sections:
            for para in section.header.paragraphs:
                if '李忠伟' in para.text:
                    count += 1
                    examples.append(f"[HEADER] {para.text.strip()[:80]}")
            for para in section.footer.paragraphs:
                if '李忠伟' in para.text:
                    count += 1
                    examples.append(f"[FOOTER] {para.text.strip()[:80]}")

        if count > 0:
            print(f"  [{count}x] {fn}: ")
            for ex in examples[:3]:
                print(f"         {ex}")
    except Exception as e:
        pass
