# -*- coding: utf-8 -*-
"""Verify 金狼电器 company replacement in generated docs."""
import os, sys, io
os.environ["PYTHONUTF8"] = "1"
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace', line_buffering=True)
from docx import Document

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")

FILES = ["3.程序.docx", "2质量手册.docx", "检验人员培训考试.docx"]
# Old values should NOT be in output
OLD = ["宁波市海曙新艺洁具有限公司", "夏亚明", "贾宇", "龚伟", "2025-03-05"]
# New values SHOULD be in output
NEW = ["宁波市金狼电器有限公司", "史迪华", "黄金金", "陈啸", "2025-04-20"]

def get_all_text(doc):
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    texts.append(p.text)
    for section in doc.sections:
        for h_obj in [section.header, section.first_page_header, section.even_page_header]:
            if h_obj and not h_obj.is_linked_to_previous:
                for p in h_obj.paragraphs:
                    texts.append(p.text)
                for t in h_obj.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                texts.append(p.text)
        for f_obj in [section.footer, section.first_page_footer, section.even_page_footer]:
            if f_obj and not f_obj.is_linked_to_previous:
                for p in f_obj.paragraphs:
                    texts.append(p.text)
    return " ".join(texts)

for fn in FILES:
    fp = os.path.join(OUTPUT_DIR, fn)
    doc = Document(fp)
    text = get_all_text(doc)
    print(f"\n=== {fn} ===")
    found_old = [v for v in OLD if v in text]
    found_new = [v for v in NEW if v in text]
    print(f"  Old values remaining: {found_old if found_old else 'NONE (all replaced!)'}")
    print(f"  New values found: {found_new}")
