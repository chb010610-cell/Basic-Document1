# -*- coding: utf-8 -*-
import sys, os, io, json, urllib.request
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.path.insert(0, r'C:\Users\Administrator\.workbuddy\binaries\python\envs\default\Lib\site-packages')

data = {
    "company_name": "\u6e29\u5dde\u6d4b\u8bd5\u7535\u5668\u6709\u9650\u516c\u53f8",
    "address": "\u6d59\u6c1f\u7701\u6e29\u5dde\u5e02\u9e7f\u57ce\u533a\u6d4b\u8bd5\u8def88\u53f7",
    "product": "\u63d2\u5ea7\u3001\u5f00\u5173\u7b49\u7535\u5668\u914d\u4ef6",
    "english_abbr": "CS",
    "doc_date": "2026-07-01",
    "version_no": "C/0",
    "gm_name": "\u738b\u6d4b\u8bd5",
    "production_mgr": "\u674e\u751f\u4ea7",
    "mgr_rep": "\u5f20\u7ba1\u7406",
    "hr_head": "\u8d77\u4e8b\u4eba",
    "sales_head": "\u94b1\u4e1a\u52a1",
    "purchase_head": "\u5b09\u91c7\u8d2d",
    "tech_head": "\u5468\u6280\u672f",
    "raw_warehouse": "\u674e\u4ed3\u5e93",
    "finished_warehouse": "\u90d1\u6210\u54c1",
    "mechanic": "\u673a\u4fee",
    "iqc_person": "\u9648\u8fdb\u6599",
    "ipqc1_person": "\u8916\u5236\u7a0b1",
    "ipqc2_person": "\u536b\u5236\u7a0b2",
    "ipqc3_person": "\u8512\u5236\u7a0b3",
    "fqc_person": "\u6c88\u6210\u54c1",
    "supplier": "\u6d4b\u8bd5\u4f9b\u5e94\u5546\u6709\u9650\u516c\u53f8",
    "license_code": "91330300MA5TEST02"
}

url = "http://127.0.0.1:8765/api/generate"
req = urllib.request.Request(url, data=json.dumps(data).encode('utf-8'),
                              headers={'Content-Type': 'application/json; charset=utf-8'})
resp = urllib.request.urlopen(req, timeout=120)
result = json.loads(resp.read().decode('utf-8'))
print(f"Success: {result['success']}  Count: {result['count']}")
total_repl = 0
for g in result.get('generated', []):
    repl = g.get('replacements', 0)
    err = g.get('error', '')
    total_repl += repl
    if err:
        print(f"  [FAIL] {g['name']}: {err}")
    else:
        print(f"  [OK] {g['name']}: {repl} replacements")
print(f"\nTotal replacements across all docs: {total_repl}")

# Verify header replacement in 3.程序.docx
from docx import Document
doc3_path = os.path.join(r'C:\Users\Administrator\WorkBuddy\20260429103845\output', '3.程序.docx')
if os.path.exists(doc3_path):
    doc = Document(doc3_path)
    print("\n=== Verify 3.程序.docx headers ===")
    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                t = para.text.strip()
                if '温州' in t or 'CS' or '2026-07' in t or 'C/0' in t:
                    print(f"  HEADER PARA: {t[:80]}")
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    ct = cell.text.strip()
                    if ct and ('温州' in ct or '2026-07' in ct or 'C/0' in ct):
                        print(f"  HEADER TABLE CELL: {ct[:80]}")

# Verify 李忠伟 replacement in 检验人员培训考试.docx
exam_path = os.path.join(r'C:\Users\Administrator\WorkBuddy\20260429103845\output', '检验人员培训考试.docx')
if os.path.exists(exam_path):
    doc = Document(exam_path)
    print("\n=== Verify 检验人员培训考试.docx (李忠伟 check) ===")
    found_lzw = False
    found_iqc = False
    for para in doc.paragraphs:
        t = para.text.strip()
        if '李忠伟' in t:
            found_lzw = True
            print(f"  [!] Still has 李忠伟: {t[:80]}")
        if '陈进料' in t:
            found_iqc = True
            print(f"  [OK] Has 陈进料: {t[:80]}")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if '李忠伟' in t:
                    found_lzw = True
                    print(f"  [!] Still has 李忠伟 (table): {t[:80]}")
                if '陈进料' in t:
                    found_iqc = True
                    print(f"  [OK] Has 陈进料 (table): {t[:80]}")
    if not found_lzw:
        print("  [OK] 李忠伟 has been fully replaced!")
    if not found_iqc:
        print("  [!] No 陈进料 found - may need checking")
