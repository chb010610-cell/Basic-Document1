# -*- coding: utf-8 -*-
"""Scan ALL source docx files (body + header/footer) and the import.xls to build complete replacement map."""
import os, sys, json
os.environ["PYTHONUTF8"] = "1"

from docx import Document

SOURCE_DIR = os.path.join(os.path.expanduser("~"), "Desktop", "website后台", "源文件")
XLS_PATH = os.path.join(os.path.expanduser("~"), "Desktop", "导入.xls")

OUT_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "scan_all_output.txt")

def scan_docx(filepath):
    """Scan a docx file and return all text content (body + header/footer)."""
    doc = Document(filepath)
    result = {"file": os.path.basename(filepath), "body_paras": [], "body_tables": [], "headers": [], "footers": []}

    # Body paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            result["body_paras"].append(p.text.strip())

    # Body tables
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                txt = cell.text.strip()
                if txt:
                    result["body_tables"].append(f"  T{ti}R{ri}C{ci}: {txt}")

    # Headers and footers
    for si, section in enumerate(doc.sections):
        for htype_name, htype_obj in [("header", section.header), ("first_page_header", section.first_page_header),
                                       ("even_page_header", section.even_page_header),
                                       ("footer", section.footer), ("first_page_footer", section.first_page_footer),
                                       ("even_page_footer", section.even_page_footer)]:
            if htype_obj and not htype_obj.is_linked_to_previous:
                tag = f"Sec{si}_{htype_name}"
                texts = []
                for p in htype_obj.paragraphs:
                    if p.text.strip():
                        texts.append(p.text.strip())
                for t in htype_obj.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            txt = cell.text.strip()
                            if txt:
                                texts.append(f"  TABLE: {txt}")
                if texts:
                    if "header" in htype_name:
                        result["headers"].append(f"{tag}: " + " | ".join(texts))
                    else:
                        result["footers"].append(f"{tag}: " + " | ".join(texts))

    return result

def read_xls(filepath):
    """Read import.xls and return all data."""
    try:
        import xlrd
        wb = xlrd.open_workbook(filepath)
        result = []
        for si in range(wb.nsheets):
            sh = wb.sheet_by_index(si)
            result.append(f"=== Sheet: {sh.name} ({sh.nrows} rows x {sh.ncols} cols) ===")
            for ri in range(min(sh.nrows, 5)):  # first 5 rows
                row_data = []
                for ci in range(sh.ncols):
                    val = sh.cell_value(ri, ci)
                    row_data.append(str(val).strip())
                result.append(f"  Row {ri}: {row_data}")
        return "\n".join(result)
    except ImportError:
        # Try openpyxl
        try:
            import openpyxl
            wb = openpyxl.load_workbook(filepath)
            result = []
            for name in wb.sheetnames:
                sh = wb[name]
                result.append(f"=== Sheet: {name} ({sh.max_row} rows x {sh.max_column} cols) ===")
                for ri, row in enumerate(sh.iter_rows(max_row=min(sh.max_row or 1, 5), values_only=True)):
                    result.append(f"  Row {ri}: {[str(v).strip() if v else '' for v in row]}")
            return "\n".join(result)
        except ImportError as e2:
            return f"Cannot read xls: xlrd not available, openpyxl error: {e2}"
    except Exception as e:
        return f"Error reading xls: {e}"

# Main
lines = []
lines.append("=" * 80)
lines.append("FULL SCAN OF ALL SOURCE DOCX FILES")
lines.append("=" * 80)

# Scan all docx files
if os.path.isdir(SOURCE_DIR):
    for fn in sorted(os.listdir(SOURCE_DIR)):
        if fn.endswith(".docx") and not fn.startswith("~"):
            fp = os.path.join(SOURCE_DIR, fn)
            lines.append(f"\n{'='*60}")
            lines.append(f"FILE: {fn}")
            lines.append(f"{'='*60}")
            try:
                r = scan_docx(fp)
                lines.append(f"--- Body paragraphs ({len(r['body_paras'])}) ---")
                for p in r["body_paras"][:50]:  # limit
                    lines.append(f"  {p}")
                if len(r["body_paras"]) > 50:
                    lines.append(f"  ... and {len(r['body_paras'])-50} more")
                lines.append(f"--- Body tables ({len(r['body_tables'])}) ---")
                for t in r["body_tables"][:80]:
                    lines.append(t)
                if len(r["body_tables"]) > 80:
                    lines.append(f"  ... and {len(r['body_tables'])-80} more")
                lines.append(f"--- Headers ({len(r['headers'])}) ---")
                for h in r["headers"]:
                    lines.append(f"  {h}")
                lines.append(f"--- Footers ({len(r['footers'])}) ---")
                for f in r["footers"]:
                    lines.append(f"  {f}")
            except Exception as e:
                lines.append(f"ERROR: {e}")
                import traceback
                lines.append(traceback.format_exc())
else:
    lines.append(f"SOURCE_DIR not found: {SOURCE_DIR}")

# Read xls
lines.append(f"\n{'='*80}")
lines.append("IMPORT.XLS CONTENT")
lines.append(f"{'='*80}")
if os.path.exists(XLS_PATH):
    lines.append(read_xls(XLS_PATH))
else:
    lines.append(f"XLS not found: {XLS_PATH}")

# Write output
with open(OUT_FILE, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))
print(f"Scan complete. Output: {OUT_FILE}")
print(f"Total lines: {len(lines)}")
