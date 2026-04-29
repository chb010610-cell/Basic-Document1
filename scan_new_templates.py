"""Scan all new DOCX templates to find all unique text values that need replacement"""
import sys, io, os, re
from collections import Counter

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document

TEMPLATE_DIR = r"C:\Users\Administrator\Desktop\website后台"

# All 12 template files
template_files = [
    "1程序文件封面.docx",
    "3.程序.docx",
    "4管理手册.docx",
    "5受控文件清单.docx",
    "7文件发放、收回和销毁记录.docx",
    "8D报告.docx",
    "9-岗位说明书.docx",
    "10.内 管理评审.docx",
    "10质量各种任命通知(3).docx",
    "13.质量目标分解考核.docx",
    "13成品召回演练记录.docx",
    "检验人员培训考试.docx",
]

print("="*80)
print("Scanning all DOCX templates for replaceable content")
print("="*80)

all_text_parts = []

for fname in template_files:
    fpath = os.path.join(TEMPLATE_DIR, fname)
    if not os.path.exists(fpath):
        print(f"\n[NOT FOUND] {fname}")
        continue
    
    doc = Document(fpath)
    
    # Collect all paragraph text
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text.strip())
    
    # Collect all table cell text  
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text.strip())
    
    print(f"\n{'='*60}")
    print(f"[FILE] {fname}")
    print(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}, Text blocks: {len(texts)}")
    
    # Show unique non-empty text (filter out very short ones)
    seen = set()
    for t in texts:
        if len(t) >= 2 and t not in seen:
            seen.add(t)
            all_text_parts.append(t)
            # Truncate long text
            display = t[:120] + "..." if len(t) > 120 else t
            print(f"  - {display}")

print("\n\n" + "="*80)
print("SUMMARY: Total unique text blocks:", len(all_text_parts))
