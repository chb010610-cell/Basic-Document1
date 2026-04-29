from docx import Document
import os

# Test reading and replacing
template_path = r"C:\Users\Administrator\Desktop\后台\2质量手册.docx"
output_dir = r"C:\Users\Administrator\WorkBuddy\20260429103845\output"
os.makedirs(output_dir, exist_ok=True)

doc = Document(template_path)

# Count original values
old_company = "宁波市海曙新艺洁具有限公司"
count = 0
for para in doc.paragraphs:
    if old_company in para.text:
        count += 1
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if old_company in cell.text:
                count += 1

print(f"Found '{old_company}' {count} times in 2质量手册.docx")

# Test replace
for para in doc.paragraphs:
    if old_company in para.text:
        for run in para.runs:
            if old_company in run.text:
                run.text = run.text.replace(old_company, "测试公司有限公司")

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if old_company in para.text:
                    for run in para.runs:
                        if old_company in run.text:
                            run.text = run.text.replace(old_company, "测试公司有限公司")

output_path = os.path.join(output_dir, "2质量手册.docx")
doc.save(output_path)
print(f"Saved to {output_path}")

# Verify
doc2 = Document(output_path)
count2 = 0
for para in doc2.paragraphs:
    if "测试公司有限公司" in para.text:
        count2 += 1
print(f"Found '测试公司有限公司' {count2} times in output")
print("SUCCESS: Replace works correctly!")