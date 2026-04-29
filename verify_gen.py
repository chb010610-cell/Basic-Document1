# -*- coding: utf-8 -*-
"""Verify generated documents contain the new values and not old ones."""
import os, sys, io
os.environ["PYTHONUTF8"] = "1"
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace', line_buffering=True)
from docx import Document

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")

# Files to verify
FILES = ["3.程序.docx", "4管理手册.docx", "检验人员培训考试.docx", "2质量手册.docx"]
# Values that should NOT be in the output (old values)
OLD_VALUES = ["宁波市海曙新艺洁具有限公司", "B/0", "A/0", "2025-03-05", "夏亚明", "贾宇", "戴海能"]
# Values that SHOULD be in the output (new values)
NEW_VALUES = ["测试公司有限公司", "C/1", "2025-06-15", "张总", "王代表", "陈检验"]

def get_all_text(doc):
    """Get all text from a document including headers/footers."""
    texts = []
    for p in doc.paragraphs:
        texts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    texts.append(p.text)
    for section in doc.sections:
        for h_obj in [section.header, section.first_page_header, section.even_page_header]:
            if h_obj and not h_obj.is_linked_to_previous:
                for p in h_obj.paragraphs:
                    texts.append(p.text)
                for t in h_obj.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                texts.append(p.text)
    combined = " ".join(texts)
    return combined

for fn in FILES:
    fp = os.path.join(OUTPUT_DIR, fn)
    if not os.path.exists(fp):
        print(f"[SKIP] {fn} not found")
        continue
    doc = Document(fp)
    text = get_all_text(doc)

    print(f"\n=== {fn} ===")
    # Check old values
    found_old = []
    for ov in OLD_VALUES:
        if ov in text:
            found_old.append(ov)
    if found_old:
        print(f"  ❌ Still contains OLD values: {found_old}")
    else:
        print(f"  ✅ No old values found - all replaced!")

    # Check new values
    found_new = []
    for nv in NEW_VALUES:
        if nv in text:
            found_new.append(nv)
    if found_new:
        print(f"  ✅ Contains NEW values: {found_new}")
    else:
        print(f"  ⚠️ No new values found")

print("\n=== Verification Complete ===")
