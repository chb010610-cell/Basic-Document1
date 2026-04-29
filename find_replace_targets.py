"""Find the original company name and key values in templates"""
import sys, io, os, re
from collections import Counter

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

TEMPLATE_DIR = r"C:\Users\Administrator\Desktop\website后台"

template_files = [
    "1程序文件封面.docx",
    "3.程序.docx",
    "4管理手册.docx",
    "5受控文件清单.docx",
    "7文件发放、收回和销毁记录.docx",
    "8D报告.docx",
    "9-岗位说明书.docx",
    "10.内 管理评审.docx",
    "10质量各种任命通知(3).docx",
    "13.质量目标分解考核.docx",
    "13成品召回演练记录.docx",
    "检验人员培训考试.docx",
]

# Collect all text and count occurrences
all_texts = []
for fname in template_files:
    fpath = os.path.join(TEMPLATE_DIR, fname)
    if not os.path.exists(fpath):
        print(f"[NOT FOUND] {fname}")
        continue
    doc = Document(fpath)
    for para in doc.paragraphs:
        t = para.text.strip()
        if t: all_texts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t: all_texts.append(t)

# Count occurrences
counter = Counter(all_texts)

# Show most common texts (these are likely the replacement candidates)
print("="*80)
print("Most frequent text values (likely replacement targets):")
print("="*80)
for text, count in counter.most_common(100):
    if len(text) >= 2:
        display = text[:80] + "..." if len(text) > 80 else text
        # Try to decode if it's garbled
        try:
            decoded = text.encode('latin-1').decode('gbk') if any(ord(c)>127 for c in text[:10]) else text
        except:
            decoded = text
        print(f"  [{count:3d}x] {display}")

# Specifically search for company-like strings (long Chinese text that appears multiple times)
print("\n\n" + "="*80)
print("Long texts appearing multiple times (company names, addresses, intros):")
print("="*80)
for text, count in counter.items():
    if count >= 2 and len(text) >= 6:
        try:
            decoded = text.encode('latin-1').decode('gbk') if any(ord(c)>127 for c in text[:10]) else text
        except:
            decoded = text
        display = decoded[:100] + "..." if len(decoded) > 100 else decoded
        print(f"  [{count:3d}x] {display}")
