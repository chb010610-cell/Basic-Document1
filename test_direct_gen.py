# -*- coding: utf-8 -*-
"""Direct test of document generation (no HTTP) to capture errors."""
import os, sys, json
os.environ["PYTHONUTF8"] = "1"

from docx import Document

_desktop = os.path.join(os.path.expanduser("~"), "Desktop")
TEMPLATE_DIR = os.path.join(_desktop, "website后台", "源文件")
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

TEMPLATE_FILES = [
    "1程序文件封面.docx", "2质量手册.docx", "3.程序.docx", "4管理手册.docx",
    "5受控文件清单.docx", "7文件发放、收回和销毁记录.docx", "8D报告.docx",
    "9-岗位说明书.docx", "10.内 管理评审.docx", "10质量各种任命通知(3).docx",
    "13.质量目标分解考核.docx", "13成品召回演练记录.docx", "检验人员培训考试.docx"
]

ORIGINAL_VALUES = {
    "company_name":   "宁波市海曙新艺洁具有限公司",
    "address":        "浙江省宁波市海曙区横街镇桃源村",
    "product":        "拖把、刷子、扫把等日用洁具",
    "english_abbr":   "XY",
    "doc_date":       "2025-03-05",
    "version_no":     "B/0",
    "header_version": "A/0",
    "gm_name":        "夏亚明",
    "mgr_rep":        "贾宇",
    "iqc_person":     "戴海能",
}

# Simple test: just replace company_name
TEST_REPLACEMENTS = [
    ("宁波市海曙新艺洁具有限公司", "测试公司有限公司"),
    ("B/0", "C/1"),
    ("A/0", "C/1"),
    ("2025-03-05", "2025-06-15"),
    ("夏亚明", "张总"),
    ("贾宇", "王代表"),
    ("戴海能", "陈检验"),
]

def replace_in_paragraph(paragraph, old_text, new_text):
    if not old_text or old_text not in paragraph.text:
        return False
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
    # Cross-run replacement
    runs = list(paragraph.runs)
    if not runs:
        return False
    combined = ''.join(r.text for r in runs)
    if old_text not in combined:
        return False
    start = combined.find(old_text)
    end = start + len(old_text)
    new_combined = combined[:start] + new_text + combined[end:]
    char_pos = 0
    target_ri = 0
    for ri, run in enumerate(runs):
        if char_pos <= start < char_pos + len(run.text):
            target_ri = ri
            break
        char_pos += len(run.text)
    remaining_pre = combined[:start]
    for ri in range(target_ri):
        if ri < len(runs) and remaining_pre:
            take = min(len(runs[ri].text), len(remaining_pre))
            runs[ri].text = remaining_pre[:take]
            remaining_pre = remaining_pre[take:]
        else:
            if ri < len(runs): runs[ri].text = ''
    if target_ri < len(runs):
        runs[target_ri].text = new_combined
    for ri in range(target_ri + 1, len(runs)):
        runs[ri].text = ''
    return True

def replace_in_table(table, old_text, new_text):
    count = 0
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if replace_in_paragraph(para, old_text, new_text):
                    count += 1
    return count

def replace_in_document(doc, old_text, new_text):
    count = 0
    for para in doc.paragraphs:
        if replace_in_paragraph(para, old_text, new_text):
            count += 1
    for table in doc.tables:
        count += replace_in_table(table, old_text, new_text)
    for section in doc.sections:
        for h_obj in [section.header, section.first_page_header, section.even_page_header]:
            if h_obj and not h_obj.is_linked_to_previous:
                for para in h_obj.paragraphs:
                    if replace_in_paragraph(para, old_text, new_text):
                        count += 1
                for table in h_obj.tables:
                    count += replace_in_table(table, old_text, new_text)
        for f_obj in [section.footer, section.first_page_footer, section.even_page_footer]:
            if f_obj and not f_obj.is_linked_to_previous:
                for para in f_obj.paragraphs:
                    if replace_in_paragraph(para, old_text, new_text):
                        count += 1
                for table in f_obj.tables:
                    count += replace_in_table(table, old_text, new_text)
    return count

# Test each file
for fn in TEMPLATE_FILES:
    fp = os.path.join(TEMPLATE_DIR, fn)
    if not os.path.exists(fp):
        print(f"[SKIP] Not found: {fn}")
        continue
    try:
        import time
        t0 = time.time()
        doc = Document(fp)
        total = 0
        for old_t, new_t in TEST_REPLACEMENTS:
            c = replace_in_document(doc, old_t, new_t)
            total += c
        out_path = os.path.join(OUTPUT_DIR, fn)
        doc.save(out_path)
        elapsed = time.time() - t0
        print(f"[OK] {fn}: {total} replacements, {elapsed:.1f}s")
    except Exception as e:
        print(f"[FAIL] {fn}: {e}")
        import traceback; traceback.print_exc()

print("\nDone!")
