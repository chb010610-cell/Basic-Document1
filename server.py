# -*- coding: utf-8 -*-
"""
ISO质量体系文档批量生成器 v4 - 完整版
模板: website后台/源文件/ (13个DOCX)
功能: 表单填写 -> 替换docx(正文+页眉页脚) -> 预览/导出Word/PDF/打印
"""
import os, sys, io, json, shutil, zipfile, webbrowser, re
os.environ["PYTHONUTF8"] = "1"
from http.server import HTTPServer, BaseHTTPRequestHandler
from socketserver import ThreadingMixIn
from urllib.parse import urlparse, parse_qs

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace', line_buffering=True)
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace', line_buffering=True)

try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--quiet"])
    from docx import Document

# ============ Configuration ============
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_desktop = os.path.join(os.path.expanduser("~"), "Desktop")
TEMPLATE_DIR = os.path.join(_desktop, "website后台", "源文件")
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "output")
HTML_PATH = os.path.join(SCRIPT_DIR, "index.html")

# 13个源文件
TEMPLATE_FILES = [
    {"id": 1,  "name": "1程序文件封面.docx",           "file": "1程序文件封面.docx"},
    {"id": 2,  "name": "2质量手册.docx",                 "file": "2质量手册.docx"},
    {"id": 3,  "name": "3.程序.docx",                     "file": "3.程序.docx"},
    {"id": 4,  "name": "4管理手册.docx",                  "file": "4管理手册.docx"},
    {"id": 5,  "name": "5受控文件清单.docx",              "file": "5受控文件清单.docx"},
    {"id": 6,  "name": "7文件发放、收回和销毁记录.docx",   "file": "7文件发放、收回和销毁记录.docx"},
    {"id": 7,  "name": "8D报告.docx",                     "file": "8D报告.docx"},
    {"id": 8,  "name": "9-岗位说明书.docx",               "file": "9-岗位说明书.docx"},
    {"id": 9,  "name": "10.内 管理评审.docx",             "file": "10.内 管理评审.docx"},
    {"id": 10, "name": "10质量各种任命通知(3).docx",      "file": "10质量各种任命通知(3).docx"},
    {"id": 11, "name": "13.质量目标分解考核.docx",        "file": "13.质量目标分解考核.docx"},
    {"id": 12, "name": "13成品召回演练记录.docx",        "file": "13成品召回演练记录.docx"},
    {"id": 13, "name": "检验人员培训考试.docx",            "file": "检验人员培训考试.docx"},
]

# 模板中的原始值（来自导入.xls 质量导入 Row1 = 新艺洁具）
ORIGINAL_VALUES = {
    # === 基本信息 ===
    "company_name":   "宁波市海曙新艺洁具有限公司",
    "address":        "浙江省宁波市海曙区横街镇桃源村",
    "product":        "拖把、刷子、扫把等日用洁具",
    "english_abbr":   "XY",

    # === 日期和版本 ===
    "doc_date":       "2025-03-05",     # 文件日期/制修订日期（正文+页眉）
    "version_no":     "B/0",            # 正文版本号（如管理手册）
    "header_version": "A/0",            # 页眉程序文件版本号

    # === 管理层人员 ===
    "gm_name":         "夏亚明",         # 总经理
    "production_mgr":  "龚伟",          # 生产部经理
    "mgr_rep":         "贾宇",          # 管理者代表
    "hr_head":         "于静洁",        # 人事部
    "sales_head":      "项万青",        # 业务部
    "purchase_head":   "阮璐",          # 采购部
    "tech_head":       "王波",          # 技术部
    "raw_warehouse":   "张丽容",        # 原材料仓库
    "finished_warehouse":"俄底只桑",    # 成品仓库
    "mechanic":        "夏方强",        # 机修

    # === 检验人员 ===
    "iqc_person":     "戴海能",          # 进料检验
    "ipqc1_person":   "解欧红",          # 制程检验1
    "ipqc2_person":   "五金",            # 制程检验2
    "ipqc3_person":   "李丹",            # 制程检验3
    "fqc_person":     "贾宇",            # 成品检验

    # === 其他信息 ===
    "supplier":        "浙江旭华包装有限公司",  # 进料供应商
    "license_code":    "91330212X10144809X",   # 营业执照编码

    # === 公司简介 ===
    "company_intro":   "",  # 从xls读取，很长

    # === 工序 ===
    "process_raw":     "原料",
    "process_1":       "五金",
    "process_2":       "注塑☆",
    "process_3":       "纺织",
    "process_4":       "裁剪",
    "process_5":       "缝纫",
    "process_6":       "装配",
    "process_7":       "包装",
    "special_process": "注塑☆",

    # === 抽样标准 ===
    "iqc_sampling":    "",
    "ipqc_sampling":   "",
    "fqc_sampling":    "",

    # === 质量方针目标 ===
    "quality_goals":   "",
    "quality_policy":  "",

    # === 时间字段 ===
    "program_time":    "2025-03-05",     # 程序时间
    "file_time":       "2025-03-05",     # 文件时间
    "audit_plan_time": "2025年11月13日", # 内审计划时间
    "audit_time":      "2025年11月20日", # 内审时间
    "audit_improve":   "2025年12月5日",  # 内审改进时间
    "review_time":     "2025年12月8日",  # 管理评审时间
    "audit_time2":     "2025年12月15日", # 第二次内审时间

    # 外发工序
    "outsource_process": "无",
}


def replace_in_paragraph(paragraph, old_text, new_text):
    """Replace text in a paragraph while preserving formatting."""
    if not old_text or old_text not in paragraph.text:
        return False
    full_text = paragraph.text
    new_full_text = full_text.replace(old_text, new_text)
    if full_text == new_full_text:
        return False
    replaced = False
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            replaced = True
    if replaced:
        return True
    runs = list(paragraph.runs)
    if not runs:
        return False
    combined = ''.join(r.text for r in runs)
    start = combined.find(old_text)
    if start == -1:
        return False
    end = start + len(old_text)
    new_combined = combined[:start] + new_text + combined[end:]
    char_pos = 0
    target_ri = 0
    for ri, run in enumerate(runs):
        if char_pos <= start < char_pos + len(run.text):
            target_ri = ri
            break
        char_pos += len(run.text)
    remaining_pre = combined[:start]
    for ri in range(target_ri):
        if ri < len(runs) and remaining_pre:
            take = min(len(runs[ri].text), len(remaining_pre))
            runs[ri].text = remaining_pre[:take]
            remaining_pre = remaining_pre[take:]
        else:
            if ri < len(runs): runs[ri].text = ''
    if target_ri < len(runs):
        runs[target_ri].text = new_combined
    for ri in range(target_ri + 1, len(runs)):
        runs[ri].text = ''
    return True


def replace_in_table(table, old_text, new_text):
    """Replace text in all cells of a table."""
    count = 0
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if replace_in_paragraph(para, old_text, new_text):
                    count += 1
    return count


def replace_in_document(doc, old_text, new_text):
    """Replace text throughout a document (body + ALL headers/footers)."""
    count = 0
    # 正文段落和表格
    for para in doc.paragraphs:
        if replace_in_paragraph(para, old_text, new_text):
            count += 1
    for table in doc.tables:
        count += replace_in_table(table, old_text, new_text)

    # 页眉页脚 - 所有section的所有header/footer类型
    for section in doc.sections:
        header_objs = [
            ("header", section.header),
            ("first_page_header", section.first_page_header),
            ("even_page_header", section.even_page_header),
        ]
        footer_objs = [
            ("footer", section.footer),
            ("first_page_footer", section.first_page_footer),
            ("even_page_footer", section.even_page_footer),
        ]
        for _name, h_obj in header_objs:
            if h_obj and not h_obj.is_linked_to_previous:
                for para in h_obj.paragraphs:
                    if replace_in_paragraph(para, old_text, new_text):
                        count += 1
                for table in h_obj.tables:
                    count += replace_in_table(table, old_text, new_text)
        for _name, f_obj in footer_objs:
            if f_obj and not f_obj.is_linked_to_previous:
                for para in f_obj.paragraphs:
                    if replace_in_paragraph(para, old_text, new_text):
                        count += 1
                for table in f_obj.tables:
                    count += replace_in_table(table, old_text, new_text)
    return count


def build_replace_map(form_data):
    """Build the complete find-replace mapping from form data."""
    replacements = []
    ov = ORIGINAL_VALUES

    def add(old, new):
        if old and new and str(old).strip() != str(new).strip():
            replacements.append((str(old), str(new)))

    # 基本信息
    add(ov["company_name"], form_data.get("company_name", ""))
    add(ov["address"], form_data.get("address", ""))
    add(ov["product"], form_data.get("product", ""))
    add(ov["english_abbr"], form_data.get("english_abbr", ""))

    # 日期
    new_date = form_data.get("doc_date", "").strip()
    if ov["doc_date"] and new_date and ov["doc_date"] != new_date:
        add(ov["doc_date"], new_date)
        # 程序时间 和 文件时间 也替换为相同日期
        if ov["program_time"] and ov["program_time"] != new_date:
            add(ov["program_time"], new_date)
        if ov["file_time"] and ov["file_time"] != new_date:
            add(ov["file_time"], new_date)

    # 版本号 - 同时替换 B/0 和 A/0
    new_ver = form_data.get("version_no", "").strip()
    if ov["version_no"] and new_ver and ov["version_no"] != new_ver:
        add(ov["version_no"], new_ver)       # 正文 B/0
    if ov["header_version"] and new_ver and ov["header_version"] != new_ver:
        add(ov["header_version"], new_ver)   # 页眉 A/0

    # 管理层人员
    add(ov["gm_name"], form_data.get("gm_name", ""))
    add(ov["production_mgr"], form_data.get("production_mgr", ""))
    add(ov["mgr_rep"], form_data.get("mgr_rep", ""))
    add(ov["hr_head"], form_data.get("hr_head", ""))
    add(ov["sales_head"], form_data.get("sales_head", ""))
    add(ov["purchase_head"], form_data.get("purchase_head", ""))
    add(ov["tech_head"], form_data.get("tech_head", ""))
    add(ov["raw_warehouse"], form_data.get("raw_warehouse", ""))
    add(ov["finished_warehouse"], form_data.get("finished_warehouse", ""))
    add(ov["mechanic"], form_data.get("mechanic", ""))

    # 检验人员
    add(ov["iqc_person"], form_data.get("iqc_person", ""))
    add(ov["ipqc1_person"], form_data.get("ipqc1_person", ""))
    add(ov["ipqc2_person"], form_data.get("ipqc2_person", ""))
    add(ov["ipqc3_person"], form_data.get("ipqc3_person", ""))
    add(ov["fqc_person"], form_data.get("fqc_person", ""))

    # 其他
    add(ov["supplier"], form_data.get("supplier", ""))
    add(ov["license_code"], form_data.get("license_code", ""))

    # 工序
    add(ov["process_raw"], form_data.get("process_raw", ""))
    add(ov["process_1"], form_data.get("process_1", ""))
    add(ov["process_2"], form_data.get("process_2", ""))
    add(ov["process_3"], form_data.get("process_3", ""))
    add(ov["process_4"], form_data.get("process_4", ""))
    add(ov["process_5"], form_data.get("process_5", ""))
    add(ov["process_6"], form_data.get("process_6", ""))
    add(ov["process_7"], form_data.get("process_7", ""))
    add(ov["special_process"], form_data.get("special_process", ""))

    # 外发工序
    add(ov["outsource_process"], form_data.get("outsource_process", ""))

    # 抽样标准
    add(ov["iqc_sampling"], form_data.get("iqc_sampling", ""))
    add(ov["ipqc_sampling"], form_data.get("ipqc_sampling", ""))
    add(ov["fqc_sampling"], form_data.get("fqc_sampling", ""))

    # 质量方针目标
    add(ov["quality_goals"], form_data.get("quality_goals", ""))
    add(ov["quality_policy"], form_data.get("quality_policy", ""))

    # 公司简介
    add(ov["company_intro"], form_data.get("company_intro", ""))

    # 时间字段
    add(ov["audit_plan_time"], form_data.get("audit_plan_time", ""))
    add(ov["audit_time"], form_data.get("audit_time", ""))
    add(ov["audit_improve"], form_data.get("audit_improve", ""))
    add(ov["review_time"], form_data.get("review_time", ""))
    add(ov["audit_time2"], form_data.get("audit_time2", ""))

    return replacements


def generate_documents(form_data):
    """Generate all documents by applying replacements to templates."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    replacements = build_replace_map(form_data)
    generated = []

    print(f"[INFO] Starting generation with {len(replacements)} replacement pairs")
    for pair in replacements:
        print(f"  '{pair[0]}' => '{pair[1]}'")

    for tmpl in TEMPLATE_FILES:
        template_path = os.path.join(TEMPLATE_DIR, tmpl["file"])
        if not os.path.exists(template_path):
            print(f"[WARN] Template not found: {template_path}")
            continue

        try:
            doc = Document(template_path)
            total_replacements = 0
            for old_text, new_text in replacements:
                count = replace_in_document(doc, old_text, new_text)
                total_replacements += count

            output_path = os.path.join(OUTPUT_DIR, tmpl["file"])
            doc.save(output_path)

            html_path = os.path.join(OUTPUT_DIR, f"preview_{tmpl['id']}.html")
            generate_html_preview(doc, html_path, tmpl["name"])

            print(f"[OK] {tmpl['name']}: {total_replacements} replacements")
            generated.append({"id": tmpl["id"], "name": tmpl["name"], "replacements": total_replacements})

        except Exception as e:
            print(f"[FAIL] Error processing {tmpl['name']}: {e}")
            import traceback; traceback.print_exc()
            generated.append({"id": tmpl["id"], "name": tmpl["name"], "error": str(e)})

    return generated


def escape_html(text):
    return (str(text).replace("&", "&amp;").replace("<", "&lt;")
            .replace(">", "&gt;").replace('"', "&quot;").replace("\n", "<br>"))


def generate_html_preview(doc, output_path, doc_name):
    """Generate HTML preview from a docx document."""
    lines = []
    lines.append('<!DOCTYPE html><html lang="zh-CN"><head><meta charset="UTF-8">')
    lines.append(f'<title>{escape_html(doc_name)} - 预览</title>')
    lines.append('<style>')
    lines.append('body{font-family:"Microsoft YaHei",sans-serif;margin:20px;background:#f5f5f5;}')
    lines.append('.page{background:#fff;max-width:210mm;margin:10px auto;padding:30mm 25mm;box-shadow:0 2px 8px rgba(0,0,0,.1);min-height:297mm;}')
    lines.append('table{border-collapse:collapse;width:100%;margin:8px 0;font-size:12px;}')
    lines.append('td,th{border:1px solid #ccc;padding:4px 6px;vertical-align:text-top;}')
    lines.append('h1,h2,h3{color:#333;}p{text-align:justify;line-height:1.8;}')
    lines.append('th{background:#eef2ff;font-weight:bold;}')
    lines.append('.highlight{color:#c00;font-weight:bold;}')
    lines.append('</style></head><body>')

    # Body content
    lines.append(f'<div class="page"><h2 style="text-align:center">{escape_html(doc_name)}</h2><hr>')
    for p in doc.paragraphs:
        txt = escape_html(p.text)
        if txt.strip():
            style = ' class="highlight"' if any(kw in p.text for kw in ['公司','日期','版本','检验']) else ''
            lines.append(f'<p{style}>{txt}</p>')

    # Tables
    for ti, table in enumerate(doc.tables):
        lines.append(f'<table><caption>表格 {ti+1}</caption>')
        for ri, row in enumerate(table.rows):
            tag = "th" if ri == 0 else "td"
            cells = "".join(f'<{tag}>{escape_html(c.text)}</{tag}>' for c in row.cells)
            lines.append(f'<tr>{cells}</tr>')
        lines.append('</table>')

    lines.append('</div></body></html>')
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def create_zip_download():
    """Create ZIP of all output DOCX files."""
    zip_path = os.path.join(OUTPUT_DIR, "ISO_documents.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for tmpl in TEMPLATE_FILES:
            fp = os.path.join(OUTPUT_DIR, tmpl["file"])
            if os.path.exists(fp):
                zf.write(fp, tmpl["file"])
    return zip_path


# ============ HTTP Request Handler ============
class RequestHandler(BaseHTTPRequestHandler):

    def log_message(self, format, *args):
        pass  # suppress default logging

    def send_json(self, data):
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(200)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Content-Length", len(body))
        self.end_headers()
        self.wfile.write(body)

    def send_file_resp(self, path, ct, filename=None):
        with open(path, "rb") as f:
            data = f.read()
        self.send_response(200)
        self.send_header("Content-Type", ct)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Content-Length", len(data))
        if filename:
            self.send_header('Content-Disposition', f'attachment; filename="{filename}"')
        self.end_headers()
        self.wfile.write(data)

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def do_GET(self):
        parsed = urlparse(self.path)
        path = parsed.path

        if path in ("/", "/index.html"):
            self.send_file_resp(HTML_PATH, "text/html; charset=utf-8")
        elif path.startswith("/preview/"):
            doc_id = path.split("/")[-1]
            hp = os.path.join(OUTPUT_DIR, f"preview_{doc_id}.html")
            if os.path.exists(hp):
                self.send_file_resp(hp, "text/html; charset=utf-8")
            else:
                self.send_error(404, "Preview not found")
        elif path.startswith("/output/"):
            fn = path.split("/")[-1]
            fp = os.path.join(OUTPUT_DIR, fn)
            if os.path.exists(fp):
                ct = "text/html" if fn.endswith(".html") else "application/octet-stream"
                self.send_file_resp(fp, ct, fn)
            else:
                self.send_error(404, "File not found")
        elif path == "/download/word":
            zp = create_zip_download()
            self.send_file_resp(zp, "application/zip", "ISO_quality_documents.zip")
        else:
            self.send_error(404, "Not found")

    def do_POST(self):
        parsed = urlparse(self.path)
        path = parsed.path
        cl = int(self.headers.get("Content-Length", 0))
        body = self.rfile.read(cl).decode("utf-8") if cl else "{}"
        try:
            data = json.loads(body)
        except:
            data = {}

        if path == "/api/generate":
            print(f"[API] /api/generate received: company={data.get('company_name','')}, date={data.get('doc_date','')}, ver={data.get('version_no','')}")
            gen = generate_documents(data)
            success_count = len([g for g in gen if "error" not in g])
            total_repl = sum(g.get("replacements", 0) for g in gen if "error" not in g)
            self.send_json({
                "success": True,
                "count": success_count,
                "total_replacements": total_repl,
                "generated": gen,
            })
        else:
            self.send_error(404, "Not found")


class ThreadedHTTPServer(ThreadingMixIn, HTTPServer):
    daemon_threads = True


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Verify templates
    if not os.path.isdir(TEMPLATE_DIR):
        print(f"[ERROR] Template dir not found: {TEMPLATE_DIR}")
        return
    tmpl_files = [f for f in os.listdir(TEMPLATE_DIR) if f.endswith(".docx")]
    print(f"[INFO] Template dir: {TEMPLATE_DIR}")
    print(f"[INFO] Found {len(tmpl_files)} .docx files")

    port = 8765
    server = ThreadedHTTPServer(("127.0.0.1", port), RequestHandler)
    print(f"[START] ISO Document Generator v4 Started!")
    print(f"[INFO] http://127.0.0.1:{port}")
    webbrowser.open(f"http://127.0.0.1:{port}")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\n[STOP] Server stopped.")
        server.server_close()


if __name__ == "__main__":
    main()
